import os
from docx import Document
from docx.document import Document as Document_type
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from rich.progress import track


limpar_console = lambda : os.system('clear' if os.name != "nt" else "clear")

meses = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Decembro"
}


gender_flexible_words = [w[:-1] for w in['solteirx', 'divorciadx', 'viúvx', 'casadx', 'brasileirx', 'aposentadx', 'inscritx', 'domiciliadx']]



def converte_palavra_genero(palavra: str, genero: str) -> str:
    return palavra[:-1] + ('o' if genero == "HOMEM" else 'a')
    

def formata_paragrafos(documento: Document_type) -> Document_type:
    # adiciona negrito a primeira palavra quando tiver classula
    in_ceter = True

    for i, paragafo in enumerate(documento.paragraphs):

        paragafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if in_ceter else WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        if i == 0: # primeiro paragafo
            text_run = paragafo.runs[0]
            text_run.underline = True
            text_run.bold = True
            paragafo.runs[0] = text_run
            in_ceter = False
            continue

        elif paragafo.text.startswith("São Luís, "):  # para onde tem a data do documento
            in_ceter = True
            paragafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if in_ceter else WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            continue

        elif paragafo.text.startswith('Testemunha'): # area de Testemunha
            in_ceter = False
            paragafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if in_ceter else WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            continue

        if paragafo.runs:  # estilização
            first_run = paragafo.runs[0]
            
            if first_run.text:
                words = first_run.text.split(sep=":", maxsplit=1)
                first_word = words[0] + ":"if len(words) > 1 else words[0]
                remaing_text = words[1] if len(words) > 1 else ""
                first_run.bold = True if len(words) > 1 else False
                first_run.text = ''
                first_run.add_text(first_word)
                paragafo.add_run(remaing_text)
                paragafo.runs[0] = first_run

    return documento


def criar_documento(template_path:str, referencias: dict) -> Document_type:
    documento = Document(template_path)
    # modifica template
    for paragafo in track(documento.paragraphs, description="Preechendo contrato..."):
        for codigo in referencias:
            paragafo.text = paragafo.text.replace(codigo, referencias[codigo])

    documento = formata_paragrafos(documento)
    return documento


def salvar_documento(documento: Document_type, nome_arquivo: str) -> str:
    DIR_NAME = 'contratos'  # diretorio em que serao salvos os contratos criados

    if not nome_arquivo.endswith((".docx", ".doc")):
        nome_arquivo += ".docx"

    if not os.path.exists(DIR_NAME):
        os.makedirs(DIR_NAME)
        print("diretorio 'contratos' criado!")

    FULLPATH = DIR_NAME + '/' + nome_arquivo
    documento.save(FULLPATH)
    return FULLPATH
